import io
import os
import shutil
import asyncio
from datetime import datetime
from typing import Optional, cast
from uuid import uuid4

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import Column, DateTime, Integer, String, Text, JSON
from sqlalchemy.orm import Session
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.styles.style import _ParagraphStyle

from app.db.database import Base, SessionLocal, get_db
from app.models.review import Review  # 🔥 REAL-TIME MODEL
from app.services.review_worker import process_review


# ─────────────────────────── DATABASE MODEL (REPORT STORAGE) ───────────────────────────

class ReviewRecord(Base):
    __tablename__ = "reviews_history"  # ✅ avoid conflict

    id = Column(String, primary_key=True, default=lambda: str(uuid4()))
    client_name = Column(String, nullable=False, index=True)
    filename = Column(String, nullable=False)

    practice = Column(String, default="")
    reviewer = Column(String, default="")
    mode = Column(String, default="full")

    date = Column(DateTime, default=datetime.utcnow, index=True)

    score = Column(Integer, nullable=False)
    risk_rating = Column(String, nullable=False)

    issue_count = Column(Integer, default=0)
    high_count = Column(Integer, default=0)
    medium_count = Column(Integer, default=0)
    low_count = Column(Integer, default=0)

    summary_headline = Column(Text, default="")
    summary_key_findings = Column(JSON, default=list)
    client_impact = Column(Text, default="")
    executive_summary = Column(Text, default="")

    issues = Column(JSON, default=list)
    plan_steps = Column(JSON, default=list)
    plan_priority = Column(String, default="")

    created_at = Column(DateTime, default=datetime.utcnow)


# ─────────────────────────── SCHEMAS ───────────────────────────

class SaveReviewRequest(BaseModel):
    client_name: str
    filename: str
    score: int
    risk_rating: str

    issue_count: int
    high_count: int
    medium_count: int
    low_count: int

    summary_headline: str = ""
    summary_key_findings: list[str] = []
    client_impact: str = ""
    executive_summary: str = ""

    issues: list[dict] = []
    plan_steps: list[dict] = []
    plan_priority: str = ""


# ─────────────────────────── ROUTER ───────────────────────────

router = APIRouter(prefix="/reviews", tags=["reviews"])

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)


# ============================================================
# 🔥 1. REAL-TIME UPLOAD + ANALYSIS (NEW)
# ============================================================

@router.post("/upload")
async def upload_and_analyze(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    review_id = str(uuid4())
    file_path = f"{UPLOAD_DIR}/{review_id}.pdf"

    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # 🔥 create real-time review record
    review = Review(
        id=review_id,
        filename=file.filename,
        status="processing"
    )

    db.add(review)
    db.commit()

    # 🚀 async analysis
    asyncio.create_task(process_review(review_id, file_path, db))

    return {
        "review_id": review_id,
        "status": "processing"
    }


# ============================================================
# 🔥 2. FETCH LIVE ANALYSIS (fallback API)
# ============================================================

@router.get("/{review_id}/findings")
def get_findings(review_id: str, db: Session = Depends(get_db)):
    review = db.query(Review).filter(Review.id == review_id).first()

    if not review:
        raise HTTPException(status_code=404, detail="Review not found")

    return {
        "status": review.status,
        "severity": review.severity,
        "findings": review.findings or []
    }


# ============================================================
# 📊 3. HISTORY (unchanged)
# ============================================================

@router.get("/history")
def list_reviews(limit: int = 100, offset: int = 0, db: Session = Depends(get_db)):
    records = (
        db.query(ReviewRecord)
        .order_by(ReviewRecord.date.desc())
        .offset(offset)
        .limit(limit)
        .all()
    )
    return records


@router.get("/{review_id}")
def get_review(review_id: str, db: Session = Depends(get_db)):
    record = db.query(ReviewRecord).filter(ReviewRecord.id == review_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="Review not found")
    return record


# ============================================================
# 📄 4. EXPORT (unchanged)
# ============================================================

@router.get("/{review_id}/export")
def export_review(review_id: str, db: Session = Depends(get_db)):
    record = db.query(ReviewRecord).filter(ReviewRecord.id == review_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="Review not found")

    doc = DocxDocument()
    doc.add_heading("AstuteIQ Report", 0)
    doc.add_paragraph(f"Client: {record.client_name}")
    doc.add_paragraph(f"Score: {record.score}")
    doc.add_paragraph(f"Risk: {record.risk_rating}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="report.docx"'},
    )


# ============================================================
# 💾 5. SAVE FINAL REPORT
# ============================================================

@router.post("", status_code=201)
def save_review(req: SaveReviewRequest, db: Session = Depends(get_db)):
    record = ReviewRecord(**req.dict())
    db.add(record)
    db.commit()
    db.refresh(record)
    return {"id": record.id}


# ============================================================
# 🗑 DELETE
# ============================================================

@router.delete("/{review_id}", status_code=204)
def delete_review(review_id: str, db: Session = Depends(get_db)):
    record = db.query(ReviewRecord).filter(ReviewRecord.id == review_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="Review not found")

    db.delete(record)
    db.commit()